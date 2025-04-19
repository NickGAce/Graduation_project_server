from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, join, func, alias
from app.database import get_async_session
from app.room.models import rooms, blocks, floors
from app.residents.models import residents
from docx import Document
from fastapi.responses import FileResponse


router = APIRouter(
    prefix="/management",
    tags=["Management"]
)

@router.get("/residents/{resident_id}/check-in-document")
async def create_check_in_document(resident_id: int, session: AsyncSession = Depends(get_async_session)):
    async with session.begin():
        # Подготовка запроса для загрузки информации о жителе и его текущей комнате
        stmt = select(
            residents.c.full_name,
            residents.c.date_of_check_in,
            residents.c.date_of_check_out,
            rooms.c.room_number,
            blocks.c.block_name,
            floors.c.floor_number
        ).select_from(
            residents
            .join(rooms, rooms.c.id == residents.c.room_id)
            .join(blocks, blocks.c.id == rooms.c.block_id)
            .join(floors, floors.c.id == blocks.c.floor_id)
        ).where(residents.c.id == resident_id)

        result = await session.execute(stmt)
        data = result.first()

        if not data:
            raise HTTPException(status_code=404, detail="Resident not found")

        # Создание документа
        doc = Document()
        doc.add_heading('Check-In Notification', level=1)
        paragraph = doc.add_paragraph()
        paragraph.add_run('Resident Name: ').bold = True
        paragraph.add_run(f'{data.full_name}\n')
        paragraph.add_run('Room Number: ').bold = True
        paragraph.add_run(f'{data.room_number} (Block: {data.block_name}, Floor: {data.floor_number})\n')
        paragraph.add_run('Date of Check-In: ').bold = True
        paragraph.add_run(f'{data.date_of_check_in}\n')
        paragraph.add_run('Date of Check-Out: ').bold = True
        paragraph.add_run(f'{data.date_of_check_out}\n')

        # Сохранение документа
        file_path = f'check_in_notice_{resident_id}.docx'
        doc.save(file_path)

        return FileResponse(path=file_path, filename=file_path)

@router.get("/residents/{resident_id}/relocation-document")
async def create_relocation_document(resident_id: int, old_room_id: int, session: AsyncSession = Depends(get_async_session)):
    async with session.begin():
        # Создание псевдонимов для таблиц для использования в запросе
        old_rooms = alias(rooms)
        old_blocks = alias(blocks)
        old_floors = alias(floors)

        # Подготовка запроса для загрузки информации о жителе и его текущей и старой комнате
        resident_info = select(
            residents.c.full_name,
            rooms.c.room_number.label("current_room_number"),
            blocks.c.block_name.label("current_block_name"),
            floors.c.floor_number.label("current_floor_number"),
            old_rooms.c.room_number.label("old_room_number"),
            old_blocks.c.block_name.label("old_block_name"),
            old_floors.c.floor_number.label("old_floor_number")
        ).select_from(
            residents
            .join(rooms, rooms.c.id == residents.c.room_id)
            .join(blocks, blocks.c.id == rooms.c.block_id)
            .join(floors, floors.c.id == blocks.c.floor_id)
            .join(old_rooms, old_rooms.c.id == old_room_id, isouter=True)
            .join(old_blocks, old_blocks.c.id == old_rooms.c.block_id, isouter=True)
            .join(old_floors, old_floors.c.id == old_blocks.c.floor_id, isouter=True)
        ).where(
            residents.c.id == resident_id
        )

        result = await session.execute(resident_info)
        data = result.first()

        if not data:
            raise HTTPException(status_code=404, detail="Resident or room not found")

        # Создание документа
        doc = Document()
        doc.add_heading('Notification of Relocation', level=1)
        paragraph = doc.add_paragraph()
        paragraph.add_run('Resident Name: ').bold = True
        paragraph.add_run(f'{data.full_name}\n')
        paragraph.add_run('From Room: ').bold = True
        paragraph.add_run(f'{data.old_room_number} (Block: {data.old_block_name}, Floor: {data.old_floor_number})\n')
        paragraph.add_run('To Room: ').bold = True
        paragraph.add_run(f'{data.current_room_number} (Block: {data.current_block_name}, Floor: {data.current_floor_number})\n')

        # Сохранение документа
        file_path = f'relocation_notice_{resident_id}.docx'
        doc.save(file_path)

        return FileResponse(path=file_path, filename=file_path)



@router.get("/summary/all/")
async def get_floors_summary(session: AsyncSession = Depends(get_async_session)):
    # Запрос для получения суммарного количества живущих на всех этажах
    occupancy_query = (
        select(func.sum(rooms.c.current_occupancy))
        .select_from(
            rooms.join(blocks, rooms.c.block_id == blocks.c.id)
            .join(floors, blocks.c.floor_id == floors.c.id)
        )
    )

    capacity_query = (
        select(func.sum(rooms.c.max_capacity))
        .select_from(
            rooms.join(blocks, rooms.c.block_id == blocks.c.id)
            .join(floors, blocks.c.floor_id == floors.c.id)
        )
    )

    # Выполнение запросов
    result_occupancy = await session.execute(occupancy_query)
    result_capacity = await session.execute(capacity_query)

    # Получение результатов
    total_occupancy = result_occupancy.scalar()
    total_capacity = result_capacity.scalar()

    if total_occupancy is None:
        total_occupancy = 0

    if total_capacity is None:
        total_capacity = 0

    return {
        "total_occupancy": total_occupancy,
        "total_capacity": total_capacity
    }


@router.get("/available_rooms/")
async def get_available_rooms(session: AsyncSession = Depends(get_async_session)):
    # Запрос для получения доступных комнат с информацией о номере этажа и названии блока
    query = (
        select(rooms.c.id, rooms.c.block_id, rooms.c.room_number, rooms.c.max_capacity, rooms.c.current_occupancy,
               floors.c.floor_number, blocks.c.block_name)
        .select_from(
            join(rooms, blocks, rooms.c.block_id == blocks.c.id)
            .join(floors, blocks.c.floor_id == floors.c.id)
        )
        .where(rooms.c.max_capacity > rooms.c.current_occupancy)
    )

    result = await session.execute(query)
    available_rooms = result.mappings().all()

    if not available_rooms:
        return {"status": "success", "data": [], "details": "No available rooms"}

    return {"status": "success", "data": available_rooms, "details": None}


@router.get("/rooms/{room_id}/residents")
async def get_residents_by_room_id(room_id: int, session: AsyncSession = Depends(get_async_session)):
    result = await session.execute(
        select(
            residents.c.id,
            residents.c.full_name,
            residents.c.gender,
            residents.c.role,
            residents.c.citizenship,
            residents.c.faculty,
            residents.c.group_number,
            residents.c.date_of_check_in,
            residents.c.date_of_check_out,
            residents.c.email,
            residents.c.status,
            rooms.c.room_number,
            rooms.c.max_capacity,
            rooms.c.current_occupancy,
            blocks.c.block_name,
            floors.c.floor_number
        )
        .select_from(residents)
        .join(rooms, rooms.c.id == residents.c.room_id)
        .join(blocks, blocks.c.id == rooms.c.block_id)
        .join(floors, floors.c.id == blocks.c.floor_id)
        .where(rooms.c.id == room_id)
    )
    residents_data = result.mappings().all()

    # Изменим поведение для случая отсутствия жителей, возвращая пустой массив вместо ошибки
    if not residents_data:
        return {
            "status": "success",
            "data": []
        }

    # Возвращаем данные жителей, если они есть
    return {
        "status": "success",
        "data": [
            {
                "id": resident.id,
                "full_name": resident.full_name,
                "gender": resident.gender,
                "role": resident.role,
                "citizenship": resident.citizenship,
                "faculty": resident.faculty,
                "group_number": resident.group_number,
                "date_of_check_in": resident.date_of_check_in.isoformat() if resident.date_of_check_in else None,
                "date_of_check_out": resident.date_of_check_out.isoformat() if resident.date_of_check_out else None,
                "email": resident.email,
                "status": resident.status,
                "room_number": resident.room_number,
                "max_capacity": resident.max_capacity,
                "current_occupancy": resident.current_occupancy,
                "block_name": resident.block_name,
                "floor_number": resident.floor_number

            } for resident in residents_data
        ]
    }

